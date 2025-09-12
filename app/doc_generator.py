# app/doc_generator.py

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Zmieniamy sygnaturę funkcji, aby przyjmowała podsumowanie materiałów
# app/doc_generator.py

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def save_order_as_word(order, material_summary, folder_path='order_docs'):
    doc = Document()
    
    # --- Data utworzenia w prawym górnym rogu ---
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run("Data utworzenia: " + order.created_at.strftime('%Y-%m-%d %H:%M'))
    date_run.font.size = Pt(12)
    
    # --- Nagłówek ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("Zlecenie - Szwalnia")
    header_run.font.size = Pt(24)
    header_run.bold = True
    header_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    order_code_par = doc.add_paragraph()
    order_code_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    order_code_run = order_code_par.add_run("Nr zlecenia: " + order.order_code)
    order_code_run.font.size = Pt(16)
    order_code_run.bold = True

    doc.add_paragraph()

    # --- Szczegóły zlecenia ---
    details_font_size = Pt(12)
    p = doc.add_paragraph(); p.add_run("Klient: ").bold = True; p.add_run(order.client.name).font.size = details_font_size
    p = doc.add_paragraph(); p.add_run("Opis: ").bold = True; p.add_run(order.description).font.size = details_font_size
    if order.fabric:
        p = doc.add_paragraph(); p.add_run("Tkanina: ").bold = True; p.add_run(order.fabric.name).font.size = details_font_size
    p = doc.add_paragraph(); p.add_run("Logowanie: ").bold = True; p.add_run(order.login_info or 'brak').font.size = details_font_size
    p = doc.add_paragraph(); p.add_run("Termin: ").bold = True; p.add_run(order.deadline.strftime('%Y-%m-%d')).font.size = details_font_size
    p = doc.add_paragraph(); p.add_run("Zlecający: ").bold = True; p.add_run(order.zlecajacy).font.size = details_font_size
    doc.add_paragraph()

    # --- Tabela produktów ---
    doc.add_heading("Produkty", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Produkt'; hdr_cells[1].text = 'Rozmiar'; hdr_cells[2].text = 'Ilość'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    for item in order.order_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.product.name
        row_cells[1].text = item.size
        row_cells[2].text = str(item.quantity)
    doc.add_paragraph()

    # --- POCZĄTEK POPRAWKI ---
    # Sekcja: Podsumowanie materiałów
    if material_summary:
        doc.add_heading("Podsumowanie materiałów dla zlecenia:", level=2)
        for summary_item in material_summary:
            # Odczytujemy dane ze słownika i tworzymy poprawny tekst
            name = summary_item.get('name', 'Brak nazwy')
            quantity = summary_item.get('quantity', 'Brak ilości')
            line_text = f"{name}: {quantity}"
            # Dopiero sformatowany tekst dodajemy do dokumentu
            doc.add_paragraph(line_text, style='List Bullet')
    # --- KONIEC POPRAWKI ---

    # --- Zapisywanie pliku ---
    base_dir = os.path.abspath(os.path.dirname(__file__))
    docs_dir = os.path.join(base_dir, folder_path)
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)
    filename = f"{order.order_code.replace('/', '_')}_{order.id}.docx"
    filepath = os.path.join(docs_dir, filename)

    doc.save(filepath)
    return filepath